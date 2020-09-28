#!/usr/bin/env python
# coding: utf-8

# In[1]:


import random                                          #随机模块
from fractions import Fraction                         #真分数
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH          #文本对齐方式
from docx.shared import Pt                             #磅数
from docx.oxml.ns import qn                            #中文格式
import time                                            #时间模块


# In[2]:


#createNumbers——生成nNum个数量的随机数字
def createNumbers(nNum):
    
    NumberList = []                                   #创建一个随机数列表
    
    for i in range(0,nNum):                           #循环随机数nNum位
        
        number = random.randint(0,100)                #number得到随机数
        
        NumberList.append(number)                     #append是添加 随机数添加到numberList列表
    
    return NumberList


# In[3]:


#createOperators——生成OpeNum个数量的随机符号
def createOperators(grade,opeNum):
    
    operators = ['+','-','*','/']
    
    operatorList = []                                    #创建一个随机符号列表
    
    if grade == 1 or grade == 2:
        
        for i in range(0,opeNum):                       #循环随机数opeNum位
            
            index = random.randint(0,1)                 #index得到随机符号下标
        
            operatorList.append(operators[index])       #添加至符号列表
    else:
        
            for i in range(0,opeNum):                       #循环随机数opeNum位
                
                index = random.randint(0,3)                 #index得到随机符号下标
        
                operatorList.append(operators[index])       #添加至符号列表
    
    return operatorList


# In[4]:


#createSimpleProblems——生成一定数量的简单式子,proLength为式子长度，即多少个数字，proNum为式子数量，即多少条
def createSimpleProblems(grade,degreeOfDiff):
    
    if degreeOfDiff == '1':                              
        
        proLength = random.randint(2,3)                #难度系数为1
        
    elif degreeOfDiff == '2':
        
        proLength = random.randint(4,5)                #难度系数为2
        
    elif degreeOfDiff == '3':
        
        proLength = random.randint(6,7)                #难度系数为3
        
    elif degreeOfDiff == '4':
        
        proLength = random.randint(7,8)                #难度系数为4
        
    elif degreeOfDiff == '5':
        
        proLength = random.randint(9,10)               #难度系数为5
    
    else:
        
        proLength = random.randint(2,4)                #默认难度系数        
        
    numbers = createNumbers(proLength)               #创建一个随机数列表
    
    operators = createOperators(grade,proLength-1)   #创建一个随机符号列表
    
    problem = ''                                     #四则运算式子
    
    for i in range(0,len(operators)):                #循环将随机数和随机符号做拼接
        
        problem += str(numbers[i]) + ' ' + operators[i] + ' '    
        
    problem += str(numbers[i+1])    
        
    if eval(problem) < 0:                            #如果结果为负数，则重新调用该函数，直到结果为非负
        
        createSimpleProblems(grade,degreeOfDiff)
    
    else:
        
        return problem


# In[5]:


# createProperFractionPro——生成真分数加减问题
def createProperFractionPro(num,degree):
    
    operator = ['+', '-']
    
    if degree == '1':
        
        degree = random.randint(1, 2)                #难度系数为1
    
    elif degree == '2':
        
        degree = random.randint(3, 4)                #难度系数为2
        
    elif degree == '3':
        
        degree = random.randint(5, 6)                #难度系数为3
        
    elif degree == '4':
        
        degree = random.randint(7, 8)                #难度系数为4
        
    elif degree == '5':
        
        degree = random.randint(9, 10)               #难度系数为5
    
    else:
        
        degree = random.randint(1, 3)                #默认难度系数
    
    problem = ''
    
    for i in range(degree):
        
        numerator = random.randint(1, 10)                                #随机生成分子
        
        denominator = random.randint(numerator, 10)                      #真分数要求分子小于分母
        
        operatorT = operator[random.randint(0, 1)]                       #只要求+ -
        
        problem += str(numerator) + '/' + str(denominator) + operatorT   #拼接成一个问题

    numerator = random.randint(1, 10)                                    #最后一个分数
    
    denominator = random.randint(numerator, 10)
    
    problem += str(numerator) + '/' + str(denominator)
    
    if eval(problem) < 0:                                               #如果结果为负数，则重新调用该函数，直到结果为非负
        
        problem = createProperFractionPro(num,degree)
        
    return problem


# In[6]:


#getSum——问题求和 并 判断结果正确性
def getSum(problem):
    
    Sum = eval(str(problem))
    
    Sum = Fraction('{}'.format(Sum)).limit_denominator()                  #转化为分数形式
        
    print("题目:"+ problem , "=")
        
    #print('正确的答案：',Sum)
        
    print("请输入本题的答案:")   
        
    inputSum = Fraction('{}'.format(eval(input()))).limit_denominator()   #获取用户输入的答案
        
    if Sum == inputSum:
            
        print("答案正确！")
            
    else:
            
        print("答案错误！")
            
        print('正确的答案：',Sum)
            
    return Sum,inputSum


# In[7]:


#createProblems——生成num个数量的式子
def createProblems(num,titleType,gradeNum,degreeOfDiff):
    
    trueNum = 0
    
    problemsList = [['0' for i in range(3)] for i in range(int(num))]
    
    for i in range(0,int(num)):                                         #生成num个数量式子
            
        if titleType == '1':
            
            problem = createSimpleProblems(gradeNum,degreeOfDiff)       #简单四则运算题型
            
            trueSum,inputSum = getSum(problem)
            
            if trueSum == inputSum:
                
                trueNum += 1
                
        else:
            
            problem = createProperFractionPro(num,degreeOfDiff)         #真分数题型
            
            trueSum,inputSum = getSum(problem)
            
            if trueSum == inputSum:
                
                trueNum += 1
            
        problemsList[i][0] = problem                                    #保存每一个问题、用户答案、正确答案
            
        problemsList[i][1] = inputSum
            
        problemsList[i][2] = trueSum
        
    accuracyRate = trueNum/float(num)                                   #计算刷题准确率
    
    print("本次刷题准确率达到{}%".format(accuracyRate*100))
    
    return accuracyRate,problemsList
    


# In[8]:


#saveRecordToDocx——保存用户刷题记录到本地docx文档
def saveRecordToDocx(title,gradeNum,degreeOfDiff,problemNum,problemsArray,accuracyRate):

    today = time.strftime("%Y{y}%m{m}%d{d}",time.localtime()).format(y='年',m='月',d='日')  #记录当天日期
    
    document = Document()                                                                   #创建一个空文档
    
    p0 = document.add_paragraph()                                                          #初始化建立第一个自然段
    
    run0 = p0.add_run('————————————小学生四则运算自动刷题库——————————\n     请输入您的所在年级:(1,2,3,4,5,6):%d\n     请输入本次刷题题型:1,0->(100以内正整数四则运算，真分数加减运算):%d\n     请输入您要选择的难易程度:(1,2,3,4,5->简单，较易，中等，较难，难):%d\n     请输入本次刷题目标数量:%d\n     ————————————————————————————'%(int(gradeNum),int(title),int(degreeOfDiff),int(problemNum)))
    
    run0.font.name = '仿宋_GB2312'                                                #设置字体中文
    
    run0.element.rPr.rFonts.set(qn('w:eastAsia'),u'仿宋_GB2312')                  #设置字体西文
    
    run0.font.bold = True                                                         #是否加粗
    
    run0.font.size = Pt(12)                                                        #字符尺寸大小
    
    for i in range(len(problemsArray)):
        
        document.styles['Normal'].font.name = u'宋体'                               #设置文档的基础字体中文                             
        
        document.styles['Normal'].element.rPr.rFonts.set(qn('w:eastAsia'),u'宋体')  #设置文档的基础字体西文

        pl = document.add_paragraph()                                              #初始化建立第二个自然段
    
        run1 = pl.add_run('题目：%s；你的答案：%s；正确答案：%s'%(problemsArray[i][0],problemsArray[i][1],problemsArray[i][2]))
    
        run1.font.name = '仿宋_GB2312'
    
        run1.element.rPr.rFonts.set(qn('w:eastAsia'),u'仿宋_GB2312')
    
        run1.font.size = Pt(11)
    
    p2 = document.add_paragraph()                                                  #初始化建立第三个自然段
    
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER                                      #对齐方式为居中，没有则默认左对齐
    
    run2 = p2.add_run('————————————————————————————\n本次刷题准确率达到：%.3f ！'%(float(accuracyRate)))
    
    run2.font.name = '仿宋_GB2312'                                                #设置字体中文
    
    run2.element.rPr.rFonts.set(qn('w:eastAsia'),u'仿宋_GB2312')                  #设置字体西文
    
    run2.font.bold = True                                                         #是否加粗
    
    run2.font.size = Pt(12)
    
    document.save('D:\Git-projects\\autoBankApp\%s刷题记录.docx'% today)
    


# In[9]:


#getParaments——获取用户的在读年级、刷题类型、难易程度、题目数量等参数
def getParaments():
    
    print("——————小学生四则运算自动刷题库——————")
    
    print("请输入您的所在年级:(1,2,3,4,5,6)")
    
    gradeNum = int(input())
    
    print("请输入本次刷题题型:1,0->(100以内正整数四则运算，真分数加减运算)")
    
    title = int(input())
    
    print("请输入您要选择的难易程度:(1,2,3,4,5->简单，较易，中等，较难，难)")
    
    degreeOfDiff = int(input())
    
    print("请输入本次刷题目标数量:")
    
    problemNum = int(input())
    
    return gradeNum,title,degreeOfDiff,problemNum


# In[10]:


# 主函数
def main():
    
    gradeNum,title,degreeOfDiff,problemNum = getParaments()                 #获取参数
    
    problemData = [[0 for i in range(3)] for i in range(int(problemNum))]  #创建一个问题数组
            
    accuracyRate,problemData=createProblems(problemNum,title,gradeNum,degreeOfDiff)     #进行刷题训练并返回刷题记录
    
    saveRecordToDocx(title,gradeNum,degreeOfDiff,problemNum,problemData,accuracyRate)   #保存记录到本地文件


# In[11]:


if __name__ ==  '__main__':
    main()




